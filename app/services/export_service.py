"""Export service (TZ section 2.12).

Saves analysis results in four formats — TXT, JSON, Markdown, DOCX — without
data loss (TZ section 11 acceptance criterion). Each exporter returns the path
of the written file under ``outputs/``.
"""
from __future__ import annotations

import json
from pathlib import Path

from sqlalchemy.orm import Session

from app.config import settings
from app.database.models import Document, Entity
from app.utils.logger import get_logger

log = get_logger("udip.export")

SUPPORTED = ("txt", "json", "md", "docx")


def _outfile(doc: Document, fmt: str) -> Path:
    settings.output_dir.mkdir(parents=True, exist_ok=True)
    stem = Path(doc.filename).stem
    return settings.output_dir / f"{stem}_{doc.public_id[:8]}.{fmt}"


def _pages(doc: Document) -> list:
    return sorted(doc.pages, key=lambda p: p.page_number)


def export_txt(doc: Document) -> Path:
    out = _outfile(doc, "txt")
    parts = [f"{doc.filename}\n{'=' * len(doc.filename)}\n"]
    for p in _pages(doc):
        parts.append(f"\n--- Sahifa {p.page_number} ---\n{p.text or ''}")
    out.write_text("\n".join(parts), encoding="utf-8")
    return out


def export_json(doc: Document, entities: list[Entity] | None = None) -> Path:
    out = _outfile(doc, "json")
    data = {
        "document": {
            "id": doc.public_id,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "page_count": doc.page_count,
            "status": doc.status,
            "metadata": doc.doc_metadata,
        },
        "pages": [
            {"page_number": p.page_number, "text": p.text, "layout": p.layout}
            for p in _pages(doc)
        ],
        "entities": [
            {"type": e.entity_type, "value": e.value, "page": e.page_id,
             "context": e.context}
            for e in (entities or [])
        ],
    }
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def export_markdown(doc: Document) -> Path:
    out = _outfile(doc, "md")
    lines = [f"# {doc.filename}", ""]
    for p in _pages(doc):
        lines.append(f"\n## Sahifa {p.page_number}\n")
        blocks = (p.layout or {}).get("blocks") if p.layout else None
        if blocks:
            for b in blocks:
                if b.get("type") == "heading":
                    lines.append(f"### {b.get('text', '')}")
                elif b.get("type") == "table" and b.get("rows"):
                    lines.append(_md_table(b["rows"]))
                elif b.get("text"):
                    lines.append(b["text"])
                lines.append("")
        else:
            lines.append(p.text or "")
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def _md_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    md = ["| " + " | ".join(header) + " |",
          "| " + " | ".join("---" for _ in header) + " |"]
    for r in rows[1:]:
        md.append("| " + " | ".join(r) + " |")
    return "\n".join(md)


def export_docx(doc: Document) -> Path:
    out = _outfile(doc, "docx")
    try:
        import docx
    except Exception:  # pragma: no cover - python-docx optional
        # Fall back to TXT if python-docx isn't available.
        log.warning("python-docx missing; exporting TXT instead of DOCX")
        return export_txt(doc)

    d = docx.Document()
    d.add_heading(doc.filename, level=0)
    for p in _pages(doc):
        d.add_heading(f"Sahifa {p.page_number}", level=2)
        blocks = (p.layout or {}).get("blocks") if p.layout else None
        if blocks:
            for b in blocks:
                if b.get("type") == "heading":
                    d.add_heading(b.get("text", ""), level=3)
                elif b.get("type") == "table" and b.get("rows"):
                    rows = b["rows"]
                    table = d.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                    table.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci, val in enumerate(row):
                            table.rows[ri].cells[ci].text = str(val)
                elif b.get("text"):
                    d.add_paragraph(b["text"])
        else:
            d.add_paragraph(p.text or "")
    d.save(str(out))
    return out


def export_document(db: Session, document_id: int, fmt: str) -> Path:
    """Export a document to the requested format; returns the output path."""
    fmt = fmt.lower()
    if fmt not in SUPPORTED:
        raise ValueError(f"Qo'llab-quvvatlanmaydigan format: {fmt}. Mavjud: {SUPPORTED}")
    doc = db.get(Document, document_id)
    if doc is None:
        raise ValueError(f"Document {document_id} not found")

    if fmt == "txt":
        return export_txt(doc)
    if fmt == "json":
        entities = db.query(Entity).filter(Entity.document_id == document_id).all()
        return export_json(doc, entities)
    if fmt == "md":
        return export_markdown(doc)
    return export_docx(doc)
